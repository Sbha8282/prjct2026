from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title style
title_style = doc.styles.add_style('TitleStyle', WD_STYLE_TYPE.PARAGRAPH)
title_font = title_style.font
title_font.name = 'Times New Roman'
title_font.size = Pt(20)
title_font.bold = True
title_font.all_caps = True

# Heading 1
h1_style = doc.styles['Heading 1']
h1_font = h1_style.font
h1_font.name = 'Times New Roman'
h1_font.size = Pt(14)
h1_font.bold = True

# Heading 2
h2_style = doc.styles['Heading 2']
h2_font = h2_style.font
h2_font.name = 'Times New Roman'
h2_font.size = Pt(13)
h2_font.bold = True

# Add headers and footers
section = doc.sections[0]
header = section.header
header_para = header.paragraphs[0]
header_para.text = "Home Services Web Application - Project Report"
header_para.style = doc.styles['Header']

footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = "Page "
footer_para.style = doc.styles['Footer']
# Add page number
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
footer_para._element.append(parse_xml(r'<w:fldSimple %s w:instr="PAGE"/>' % nsdecls('w')))

# Save as reference
doc.save('/workspaces/prjct2026/reference.docx')

print("Reference DOCX created with Times New Roman styles, headers, and footers.")